import streamlit as st
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import io
import json

# --- Helper to force strict margins and strip default padding ---
def format_cell_structure(cell, width_inches):
    cell.width = Inches(width_inches)
    tcPr = cell._tc.get_or_add_tcPr()
    
    tcMar = OxmlElement('w:tcMar')
    for margin_type in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{margin_type}')
        node.set(qn('w:w'), '0')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# --- Helper function for completely borderless tables ---
def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    new_borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:insideH w:val="none"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(new_borders)

# --- Helper to add a top or bottom single border line to a row ---
def add_row_horizontal_border(row, border_position="bottom"):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is not None:
            tcPr.remove(tcBorders)
        new_border = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>\n'
            f'  <w:{border_position} w:val="single" w:sz="6" w:space="0" w:color="000000"/>\n'
            f'</w:tcBorders>'
        )
        tcPr.append(new_border)

# --- Docx production engine ---
def build_docx(header_data, questions_list):
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)

    # 1. School Main Details Header Area
    if header_data.get('logo_file'):
        header_table = doc.add_table(rows=1, cols=2)
        remove_table_borders(header_table)
        format_cell_structure(header_table.cell(0, 0), 1.2)
        format_cell_structure(header_table.cell(0, 1), 5.8)
        
        logo_p = header_table.cell(0, 0).paragraphs[0]
        logo_run = logo_p.add_run()
        logo_run.add_picture(header_data['logo_file'], width=Inches(1.0))
        p = header_table.cell(0, 1).paragraphs[0]
    else:
        p = doc.add_paragraph()
        
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_run = p.add_run(header_data['school_name'])
    p_run.bold = True
    p_run.font.size = Pt(14)

    if header_data.get('logo_file'):
        p_sub = header_table.cell(0, 1).add_paragraph()
    else:
        p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = p_sub.add_run(f"{header_data['address']}\nPh. No. - {header_data['phone']}, Email - {header_data['email']}")
    sub_run.font.size = Pt(9.5)

    # 2. Student & Time Metadata Grid (Borderless layout matching original flow)
    meta_table = doc.add_table(rows=2, cols=2)
    remove_table_borders(meta_table)
    for row in meta_table.rows:
        format_cell_structure(row.cells[0], 3.5)
        format_cell_structure(row.cells[1], 3.5)
        
    p_time = meta_table.cell(0, 0).paragraphs[0]
    p_time.add_run(f"Time: {header_data['time']}")
    
    p_mm = meta_table.cell(0, 1).paragraphs[0]
    p_mm.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    mm_run = p_mm.add_run(f"M.M: {header_data['max_marks']}")
    mm_run.bold = True
    
    p_name = meta_table.cell(1, 0).paragraphs[0]
    p_name.add_run("Name: ______________________")
    
    p_roll = meta_table.cell(1, 1).paragraphs[0]
    p_roll.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_roll.add_run("Roll No. ____________")

    # 3. Assessment Banner Box (Surrounded by clean top and bottom horizontal lines)
    banner_table = doc.add_table(rows=1, cols=1)
    remove_table_borders(banner_table)
    format_cell_structure(banner_table.cell(0, 0), 7.0)
    add_row_horizontal_border(banner_table.rows[0], "top")
    add_row_horizontal_border(banner_table.rows[0], "bottom")
    
    p_title = banner_table.cell(0, 0).paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run(f"{header_data['assessment_name']}\nSUBJECT - {header_data['subject']}\nCLASS - {header_data['class_name']}")
    r_title.bold = True
    r_title.font.size = Pt(11)

    # General Instructions line block
    p_ins = doc.add_paragraph()
    p_ins.paragraph_format.space_before = Pt(6)
    p_ins.paragraph_format.space_after = Pt(12)
    p_ins.add_run("General Instructions:\n1. All the questions are compulsory.\n2. Write the answers neatly in the answer sheet.")

    # 4. Dynamic Question Processing Loop
    current_main_heading = ""
    for idx, q in enumerate(questions_list, 1):
        # Render main question headers dynamically if changed
        if q.get('main_heading') and q.get('main_heading') != current_main_heading:
            current_main_heading = q['main_heading']
            h_table = doc.add_table(rows=1, cols=2)
            remove_table_borders(h_table)
            format_cell_structure(h_table.rows[0].cells[0], 5.2)
            format_cell_structure(h_table.rows[0].cells[1], 1.8)
            
            hp1 = h_table.cell(0, 0).paragraphs[0]
            hr1 = hp1.add_run(current_main_heading)
            hr1.bold = True
            
            hp2 = h_table.cell(0, 1).paragraphs[0]
            hp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if q.get('section_marks'):
                hr2 = hp2.add_run(f"({q['section_marks']})")
                hr2.bold = True
            h_table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(8)
        
        q_type = q['type']
        
        # Base Question Content Line Grid
        q_table = doc.add_table(rows=1, cols=2)
        remove_table_borders(q_table)
        format_cell_structure(q_table.rows[0].cells[0], 0.5)
        format_cell_structure(q_table.rows[0].cells[1], 6.5)
        
        num_run = q_table.cell(0, 0).paragraphs[0].add_run(f"{q.get('sub_number', idx)}")
        num_run.bold = True
        
        q_text_p = q_table.cell(0, 1).paragraphs[0]
        q_text_p.add_run(q['text'])
        
        if q_type == "MCQ":
            opt_table = doc.add_table(rows=1, cols=5)
            remove_table_borders(opt_table)
            format_cell_structure(opt_table.rows[0].cells[0], 0.5)
            for i in range(4):
                cell = opt_table.rows[0].cells[i + 1]
                format_cell_structure(cell, 1.625)
                p_opt = cell.paragraphs[0]
                p_opt.add_run(q['options'][i])
            opt_table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Pt(6)

        elif q_type == "Match the Following":
            match_table = doc.add_table(rows=len(q['pairs']), cols=3)
            remove_table_borders(match_table)
            
            for r_idx, pair in enumerate(q['pairs']):
                cell_spacer = match_table.cell(r_idx, 0)
                cell_left = match_table.cell(r_idx, 1)
                cell_right = match_table.cell(r_idx, 2)
                
                format_cell_structure(cell_spacer, 0.5)
                format_cell_structure(cell_left, 3.25)
                format_cell_structure(cell_right, 3.25)
                
                c1 = cell_left.paragraphs[0]
                c2 = cell_right.paragraphs[0]
                
                if str(pair.get('left_type')).strip().lower() == "image" and pair.get('left_img'):
                    c1.add_run(pair.get('left_prefix', '') + " ")
                    c1.add_run().add_picture(pair['left_img'], width=Inches(1.0))
                else:
                    c1.add_run(pair.get('left_text', ''))
                    
                if str(pair.get('right_type')).strip().lower() == "image" and pair.get('right_img'):
                    c2.add_run(pair.get('right_prefix', '') + " ")
                    c2.add_run().add_picture(pair['right_img'], width=Inches(1.0))
                else:
                    c2.add_run(pair.get('right_text', ''))
            match_table.rows[-1].cells[0].paragraphs[0].paragraph_format.space_after = Pt(6)

        elif q_type == "Image/Source Based":
            if q.get('image_file'):
                img_table = doc.add_table(rows=1, cols=2)
                remove_table_borders(img_table)
                format_cell_structure(img_table.cell(0, 0), 0.5)
                format_cell_structure(img_table.cell(0, 1), 6.5)
                img_table.cell(0, 1).paragraphs[0].add_run().add_picture(q['image_file'], width=Inches(3.0))
            
            for sub_idx, sub_q in enumerate(q['sub_questions'], 1):
                sub_table = doc.add_table(rows=1, cols=3)
                remove_table_borders(sub_table)
                format_cell_structure(sub_table.rows[0].cells[0], 0.5)
                format_cell_structure(sub_table.rows[0].cells[1], 0.5)
                format_cell_structure(sub_table.rows[0].cells[2], 6.0)
                
                sub_table.cell(0, 1).paragraphs[0].add_run(f"({sub_idx})")
                sub_table.cell(0, 2).paragraphs[0].add_run(sub_q)
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            
        else:
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream

# --- Streamlit Frontend UI ---
st.set_page_config(page_title="Exam Creator Pro", layout="wide")
st.title("📝 Automatic Exam Sheet Template Generator")

if 'questions' not in st.session_state:
    st.session_state.questions = []

with st.sidebar:
    st.header("🏫 School & Header Info")
    uploaded_logo = st.file_uploader("Upload School Logo", type=["png", "jpg", "jpeg"])
    
    h_data = {
        "logo_file": uploaded_logo,
        "school_name": st.text_input("School/Institution Name", "VSI GLOBAL SR. SEC. SCHOOL"),
        "address": st.text_area("Address Line", "Sec. 5, Pratap Nagar, Behind Pratap Plaza Tonk Road, Sanganer, Jaipur (Raj.)"),
        "phone": st.text_input("Phone Number", "9309305656"),
        "email": st.text_input("Email ID", "vsiglobalschool@gmail.com"),
        "assessment_name": st.text_input("Assessment Title", "ASSESSMENT SHEET - 2026-27"),
        "subject": st.text_input("Subject", "EVS"),
        "class_name": st.text_input("Class Level", "III"),
        "time": st.text_input("Time Duration Limit", "1 HOUR"),
        "max_marks": st.text_input("Maximum Marks (M.M)", "20")
    }

# --- JSON Import Processing ---
st.subheader("📂 Bulk Import Options")
uploaded_json = st.file_uploader("Upload pre-configured JSON question bank file", type=["json"])

if uploaded_json is not None:
    try:
        imported_data = json.load(uploaded_json)
        if isinstance(imported_data, list):
            if st.button("📥 Overwrite and load questions from JSON"):
                for q in imported_data:
                    if q['type'] == "Match the Following" and 'pairs' in q:
                        for p in q['pairs']:
                            if 'left_type' not in p: p['left_type'] = "Text"
                            if 'right_type' not in p: p['right_type'] = "Text"
                            p['left_img'] = None
                            p['right_img'] = None
                st.session_state.questions = imported_data
                st.success(f"Successfully imported {len(imported_data)} questions!")
                st.rerun()
        else:
            st.error("Invalid format: The JSON root layout must be a list.")
    except Exception as e:
        st.error(f"Error reading JSON file: {e}")

st.write("---")

st.subheader("🛠️ Step 2: Build Paper Structure")
col_num, col_type, col_add = st.columns([2, 4, 2])

with col_type:
    q_type_sel = st.selectbox("Choose Question Type to Add", ["MCQ", "Fill in Blanks / Short Ans", "Match the Following", "Image/Source Based"])

with col_add:
    st.write("##")
    if st.button("➕ Add This Question Element"):
        new_q = {
            "type": q_type_sel, 
            "main_heading": "Ques.1 Tick ( ) the Correct Answer.", 
            "section_marks": "0.5 × 5 = 2.5", 
            "sub_number": "i)", 
            "text": "", 
            "marks": ""
        }
        if q_type_sel == "MCQ":
            new_q["options"] = ["(a) ", "(b) ", "(c) ", "(d) "]
        elif q_type_sel == "Match the Following":
            new_q["pairs"] = [{
                "left_type": "Text", "left_text": "i) Item", "left_prefix": "i)", "left_img": None,
                "right_type": "Text", "right_text": "a. Target", "right_prefix": "a.", "right_img": None
            }]
        elif q_type_sel == "Image/Source Based":
            new_q["image_file"] = None
            new_q["sub_questions"] = [""]
        st.session_state.questions.append(new_q)

# Render forms interactively
for idx, question in enumerate(st.session_state.questions):
    with st.expander(f"Question Element N°{idx+1} — Form Category Layout: **{question['type']}**", expanded=True):
        ch1, ch2 = st.columns(2)
        question['main_heading'] = ch1.text_input("Section Header (e.g. Ques.1...)", value=question.get('main_heading', ''), key=f"mh_{idx}")
        question['section_marks'] = ch2.text_input("Section Total Marks (e.g. 0.5 × 5 = 2.5)", value=question.get('section_marks', ''), key=f"sm_{idx}")
        
        c_num, c_q = st.columns([2, 6])
        question['sub_number'] = c_num.text_input("Question Numbering Index", value=question.get('sub_number', str(idx+1)), key=f"num_{idx}")
        question['text'] = c_q.text_area(f"Question/Instruction Text String", value=question['text'], key=f"txt_{idx}")
        
        if question['type'] == "MCQ":
            st.markdown("**Enter Option Values below:**")
            opt_cols = st.columns(4)
            for o_idx in range(4):
                question['options'][o_idx] = opt_cols[o_idx].text_input(f"Option {chr(97+o_idx)}", value=question['options'][o_idx], key=f"opt_{idx}_{o_idx}")
                
        elif question['type'] == "Match the Following":
            st.markdown("**Define Match Pairs:**")
            c_add_p, c_rem_p = st.columns([1, 1])
            if c_add_p.button("➕ Append row pair line", key=f"add_pair_{idx}"):
                question['pairs'].append({
                    "left_type": "Text", "left_text": "i) Item", "left_prefix": "i)", "left_img": None,
                    "right_type": "Text", "right_text": "a. Target", "right_prefix": "a.", "right_img": None
                })
                st.rerun()
            if c_rem_p.button("➖ Remove last row pair", key=f"rem_pair_{idx}"):
                if len(question['pairs']) > 1:
                    question['pairs'].pop()
                    st.rerun()
            
            for p_idx, pair in enumerate(question['pairs']):
                st.markdown(f"**Pair Row {p_idx+1}**")
                l_col, r_col = st.columns(2)
                
                with l_col:
                    pair['left_type'] = st.radio(f"Left Type ({p_idx+1})", ["Text", "Image"], index=0 if pair.get('left_type')=="Text" else 1, key=f"ltype_{idx}_{p_idx}")
                    pair['left_prefix'] = st.text_input(f"Left Index (e.g. i)", value=pair.get('left_prefix', ''), key=f"lpref_{idx}_{p_idx}")
                    if pair['left_type'] == "Text":
                        pair['left_text'] = st.text_input(f"Left Text String", value=pair.get('left_text', ''), key=f"ltxt_{idx}_{p_idx}")
                
                with r_col:
                    pair['right_type'] = st.radio(f"Right Type ({p_idx+1})", ["Text", "Image"], index=0 if pair.get('right_type')=="Text" else 1, key=f"rtype_{idx}_{p_idx}")
                    pair['right_prefix'] = st.text_input(f"Right Index (e.g. a.)", value=pair.get('right_prefix', ''), key=f"rpref_{idx}_{p_idx}")
                    if pair['right_type'] == "Text":
                        pair['right_text'] = st.text_input(f"Right Text String", value=pair.get('right_text', ''), key=f"rtxt_{idx}_{p_idx}")

        elif question['type'] == "Image/Source Based":
            question['image_file'] = st.file_uploader(f"Upload Image/Source Context File {idx+1}", type=["png", "jpg", "jpeg"], key=f"ctx_img_{idx}")
            st.markdown("**Sub-questions attached below context:**")
            
            c_add_s, c_rem_s = st.columns([1, 1])
            if c_add_s.button("➕ Append sub-question step line", key=f"add_sub_{idx}"):
                question['sub_questions'].append("")
                st.rerun()
            if c_rem_s.button("➖ Remove last sub-question", key=f"rem_sub_{idx}"):
                if len(question['sub_questions']) > 1:
                    question['sub_questions'].pop()
                    st.rerun()
                    
            for s_idx, sub_q in enumerate(question['sub_questions']):
                question['sub_questions'][s_idx] = st.text_input(f"Sub-question ({s_idx+1}) Text", value=sub_q, key=f"sub_{idx}_{s_idx}")

        st.write("##")
        if st.button("🗑️ Delete Entire Question Block", key=f"del_{idx}"):
            st.session_state.questions.pop(idx)
            st.rerun()

st.write("---")
if st.session_state.questions:
    backup_ready_questions = []
    for q in st.session_state.questions:
        q_copy = q.copy()
        if "image_file" in q_copy:
            q_copy["image_file"] = None
        if q_copy.get("type") == "Match the Following" and "pairs" in q_copy:
            cleaned_pairs = []
            for pair in q_copy["pairs"]:
                p_copy = pair.copy()
                if "left_img" in p_copy: p_copy["left_img"] = None
                if "right_img" in p_copy: p_copy["right_img"] = None
                cleaned_pairs.append(p_copy)
            q_copy["pairs"] = cleaned_pairs
        backup_ready_questions.append(q_copy)

    json_str = json.dumps(backup_ready_questions, indent=2)
    st.sidebar.download_button(
        label="💾 Save Current Question Stack to JSON File",
        data=json_str,
        file_name="question_backup.json",
        mime="application/json"
    )
    
    if st.button("🚀 Render and Compile Document Grid Template"):
        with st.spinner("Assembling layout structure properties safely..."):
            docx_buffer = build_docx(h_data, st.session_state.questions)
            st.success("Compilation Success! Your flawlessly aligned document is ready.")
            st.download_button(
                label="📥 Download Microsoft Word (.docx) File",
                data=docx_buffer,
                file_name=f"{h_data['subject']}_Exam_Template.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
